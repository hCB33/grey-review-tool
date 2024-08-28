#!/usr/bin/env python3
"""
Created on Wed Jul 31 14:42:16 2024

@author: Hamish Clark
@title: Grey Literature Search Tool
@desc: This python script is used to search through the gov.uk site using web 
        scraping to find any related literature to some keywords provided by 
        either the user or a generative AI model
"""
# Used for searching through web URLs and retrieving information
from bs4 import BeautifulSoup
import requests
import re

# Used to read in dataset containing all departments, public bodies, etc. (pandas is also used)
import csv

# Used for front-end design of the tool
import tkinter as tk
from tkinter import messagebox
from tkinter.scrolledtext import ScrolledText
from functools import partial
import webbrowser
import tabulate

# Used to check the input for any date provided by the user
import datetime as dt
from datetime import datetime

# Used to export results to Excel or Word
import pandas as pd
import os
import docx


# Class HyperLinkManager taken from https://stackoverflow.com/questions/76326100/how-to-add-hyperlink-to-a-tkinter-output-text
class HyperlinkManager:
    def __init__(self, text):
        self.text = text
        self.text.tag_config("hyper", foreground="blue", underline=1)
        self.text.tag_bind("hyper", "<Button-1>", self._click)
        self.reset()

    def reset(self):
        self.links = {}

    def add(self, action):
        # add an action to the manager.  returns tags to use in
        # associated text widget
        tag = "hyper-%d" % len(self.links)
        self.links[tag] = action
        return "hyper", tag

    def _click(self, event):
        for tag in self.text.tag_names(tk.CURRENT):
            if tag[:6] == "hyper-":
                self.links[tag]()
                return

headers = {
    'User-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/105.0.0.0 Safari/537.36'
}

# Returns the URL to the departments blog page
def get_blog(data):
  blog_link = 'None'
  content = data.find("ul", {"class": "gem-c-share-links__list"})
  try:
    for i in content.find_all("li"):
      if "blog" in i.find("a").text.strip().lower():
        blog_link = i.find("a")['href']
  except:
    try:
      content = data.find("div", {"class": "govuk-notification-banner__content"})
      blog_link = content.find("a")['href']
    except:
      blog_link='None'
  return blog_link

# =============================================================================
# Returns a list of all the departments info from gov.uk (Department name, URL
# to gov.uk page, and the URL to their corresponding blog page, takes about 10 
# mins to complete)
# =============================================================================
def all_deps_ukgov():
  html = requests.get("https://www.gov.uk/government/organisations", headers=headers).text
  soup = BeautifulSoup(html, 'html.parser')
  result = []
  for i in soup.find_all('li', {'class': 'organisations-list__item'}):
    links = []
    for j in i.find_all('a'):
      title = j.text.replace("\r\n", "").strip("\n")
      title_link = f"https://www.gov.uk{j['href']}"
      blog_html = requests.get(title_link, headers=headers).text
      data = BeautifulSoup(blog_html, 'html.parser')
      blog_link, only_blog = get_blog(data)
      entry = {'Title': title, 'Link': title_link, 'Blog Link': blog_link, 'Only Blog': only_blog}
      if len(links)==0:
        entry['Works with'] = []
        links.append(entry)
      elif len(links)==1:
        links[0]['Works with'].append(entry)
    if len(links[0]['Works with'])==0:
      links[0]['Works with'] = 'None'
    result.append(links[0])
  return result

# =============================================================================
# Takes a dictionary in the form of a string and converts it back to a 
# dicitonary
# =============================================================================
def format_dict(d):
  dictionary = dict()
  pairs = d.strip('{}').split(', \'')
  for i in pairs:
    pair = i.split(': ')
    dictionary[pair[0].strip('\'\'\"\"')] = pair[1].strip('\'\'\"\"')
  return dictionary

# =============================================================================
# Given a list of dictionaries as a string, converts string back to list of 
# dictionaries
# =============================================================================
def format_list_dicts(s):
  result = []
  s_sep = s[1:-1]
  s_split = s_sep.split("}, ")
  for i in range(len(s_split)):
    if s_split[i][-1]!="}":
      s_split[i]+="}"
    d = format_dict(s_split[i])
    result.append(d)
  return result

# =============================================================================
# Returns the same information as all_deps_ukgov but from a locally stored csv 
# file, making it quicker.
# =============================================================================
def all_deps_csv():
    with open('dataset.csv') as f:
      a = [{k: v for k, v in row.items()} for row in csv.DictReader(f, skipinitialspace=True)]
    for i in range(len(a)):
      if a[i]['Works with']=="[]":
        a[i]['Works with'] = 'None'
      else:
        a[i]['Works with'] = format_list_dicts(a[i]['Works with'])
    return a

# user_choice = ""
# print("Would you like to select from a list of departments from the webaite or a recent version? (Y/N)")
# print("Note: The online version will take longer to load (about 5-9 mins)")
# while True:
#     user_choice = input("> ")
#     if user_choice in ["Y", "N"]:
#         break
#     else:
#         print("ERROR: Please enter either Y or N")

# df = None
# if user_choice == "Y":
#     print("Retrieving information from web...")
#     df = all_deps_ukgov()
#     print("Done!")
# else:
#     print("Retrieving information locally...")
#     df = all_deps_csv()
#     print("Done!")

# =============================================================================
# Reads 'read_blogs.csv' and returns dataset containing necessary information
# to retrieve information from those blog pages
# =============================================================================
def read_all_blogs():
    with open('read_blogs.csv') as f:
        all_blogs = [{k: v for k, v in row.items()} for row in csv.DictReader(f, skipinitialspace=True)]
    for i in range(len(all_blogs)):
        pages = all_blogs[i]['Page'].split(" ")
        try:
            all_blogs[i]['Page Increment'] = int(all_blogs[i]['Page Increment'])
        except:
            pass
        if len(pages)>1:
            all_blogs[i]['Page'] = pages
    return all_blogs[:12]#----------REMOVE [:12] ONCE ALL BLOG DATA IS FILLED--------------------------------
#------------------------------------------------------------------------------------------------------------


try:
    df = all_deps_csv()
except:
    df = all_deps_ukgov()
    pd.DataFrame(df).to_csv('dataset.csv', index=False)

all_blogs = read_all_blogs()
blog_titles = []
for elem in all_blogs:
    blog_titles.append(elem['Title'])



#---------------------- Grey Literature Search Tool ---------------------------

# =============================================================================
# Find and returns the date in which a gov.uk page was first published
# =============================================================================
def get_og_gov_date(link):
    date = "N/A"
    try:
        html = requests.get(link, headers=headers).text
        soup = BeautifulSoup(html, 'html.parser')
        date = str(soup.find('meta', {'name': 'govuk:first-published-at'})).split("\"")[1][:10].split("-")
        date = "/".join([date[2], date[1], date[0]])
    except:
        pass
    return date

# =============================================================================
# Returns the search URL given the search criteria set by the user
# =============================================================================
def govuk_pubs_link(df, search_terms, from_date=None, end_date=None, sort_by=None):
  orgs = []
  url_search = "+".join(search_terms.split(" "))
  for i in range(len(df)):
    link = df[i]['Link']
    org = link.split("/")[-1]
    orgs.append(org)
  URL = "https://www.gov.uk/search/all?"
  for org in orgs:
    URL += f"&organisations[]={org}"
  if sort_by=="Relevance":
      URL += "&order=relevance"
  elif sort_by=="Newest First":
      URL += "&order=updated-newest"
  elif sort_by=="Newest First":
      URL += "&order=updated-oldest"
  else:
      URL += "&order=relevance"
  URL += f"&keywords={url_search}"
  if from_date!=None and end_date!=None:
    URL += f"&public_timestamp%5Bfrom%5D={from_date}&public_timestamp%5Bto%5D={end_date}"
  elif from_date!=None:
      URL += f"&public_timestamp%5Bfrom%5D={from_date}"
  elif end_date!=None:
      URL += f"&public_timestamp%5Bto%5D={end_date}"
  URL += "&page=1"
  return URL

# =============================================================================
# Returns the authors of the literature provided by the link
# =============================================================================
def get_author_deps(link):
  html = requests.get(link, headers=headers).text
  data = BeautifulSoup(html, 'html.parser')
  deps = []
  parent = data.find("body").find_all("li", {"class": "organisation-logos__logo"})
  for elem in parent:
    department = elem.text.strip()
    if department not in titles:
      words = re.split(r'(?<=[a-z])(?=[A-Z])|(?<=[A-Z])(?=[A-Z][a-z])', department)
      department = ' '.join(words)
      if department in titles:
        deps.append(department)
    else:
      deps.append(department)
  if len(deps)==0:
    try:
      parent = data.find("body").find("dd", {"class": "gem-c-metadata__definition"}).find_all("a", {"class": "govuk-link"})
    except:
      try:
        parent = data.find("body").find("div", {"class": "organisations-list"}).find_all("a", {"class": "govuk-link"})
      except:
        try:
          parent = data.find("body").find("div", {"class": "gem-c-organisation-logo brand--executive-office"}).find_all("span", {"class": "gem-c-organisation-logo__name"})
        except:
            try:
                parent = data.find("body").find("div", {"class": "govuk-!-width-one-half govuk-!-margin-top-3 responsive-bottom-margin"}).find_all("span", {"class": "gem-c-organisation-logo__name"})
            except:
                try:
                    parent = data.find("body").find("div", {"class": "gem-c-organisation-logo brand--attorney-generals-office"}).find_all("span", {"class": "gem-c-organisation-logo__name"})
                except:
                    return deps
    for elem in parent:
      department = elem.text.strip()
      deps.append(department)
  return deps

# =============================================================================
# Returns a list containing all the related literature from the html data given
# =============================================================================
def get_list_govuk(df, data):
  result = []
  parent = data.find("body").find("div", {"class": "finder-results js-finder-results"})
  text = list(parent.descendants)[1]
  for i in text.find_all("li", {"class": "gem-c-document-list__item"}):
    title = i.find("a",{"class": "govuk-link"}).text.strip()
    title_link = i.find("a",{"class": "govuk-link"})['href']
    try:
      desc = i.find("p", {"class": "gem-c-document-list__item-description"}).text.strip()
    except:
      desc = 'None'
    updated = i.find("ul", {"class": "gem-c-document-list__item-metadata"}).text.strip()[9:]
    authors = get_author_deps(f"https://www.gov.uk{title_link}")
    if len(authors)==0:
      authors = "N/A"
      for elem in df:
          if elem['Link']==f"https://www.gov.uk{title_link}":
              authors = elem['Title']
              break
    result.append({
        "Title": title,
        "URL": f"https://www.gov.uk{title_link}",
        "Departments, Agencies, and Public bodies": authors,
        "Abstract": desc,
        "Last Updated": updated,
        "Date Published": get_og_gov_date(f"https://www.gov.uk{title_link}"),
    })
  return result


# =============================================================================
# Retrieves total number of results
# =============================================================================
def get_total_results(link, selected_blogs=None):
    html = requests.get(link, headers=headers).text
    data = BeautifulSoup(html, 'html.parser')
    try:
        no_results = data.find("div", {"class": "result-info__header"}).text.strip().split(" ")
        if int(no_results[0].replace(',',''))==0:
          return None
        if selected_blogs!=None:
            no_results = int(no_results[0].replace(',',''))
            for elem in selected_blogs:
                number = 0
                if elem[0]['Number']=="MANUAL":
                    print("is manual")
                    number = find_blog_number(elem)
                else:
                    blog = elem[0]
                    link = elem[1]
                    try:
                        conn = requests.get(link, headers=headers)
                        html = conn.text
                        soup = BeautifulSoup(html, 'html.parser')
                        b_num_results = blog['Number'].split(" ")
                        if len(b_num_results)>3:
                            temp = " ".join(b_num_results[1:-1])
                            b_num_results = [b_num_results[0], temp, b_num_results[-1]]
                        i = soup.find(b_num_results[0], {'class': b_num_results[1]})
                        number = i.text.strip().split(" ")[int(b_num_results[2])].strip()
                        number = int(number.replace(',',''))
                    except:
                        number = 0
                if number==0:
                    selected_blogs.remove(elem)
                no_results += number
        else:        
            no_results = int(no_results[0].replace(',',''))     
        return no_results, selected_blogs
    except:
      return None, None

# =============================================================================
# Returns a list of results that does not contain any duplicates
# =============================================================================
def remove_dupes(result):
    seen = set()
    new_l = []
    for d in result:
        if type(d['Departments, Agencies, and Public bodies'])==list:
            d['Departments, Agencies, and Public bodies'] = tuple(d['Departments, Agencies, and Public bodies'])
        t = tuple(d.items())
        if t not in seen:
            seen.add(t)
            new_l.append(d)
    for i in range(len(new_l)):
        if type(new_l[i]['Departments, Agencies, and Public bodies'])==tuple:
            new_l[i]['Departments, Agencies, and Public bodies'] = list(new_l[i]['Departments, Agencies, and Public bodies'])
        
    return new_l


# =============================================================================
# Returns a list containing all the related literature given the max number of
# results provided by the user
# =============================================================================
def get_pubs(df, link, max_results, blogs=None):
  html = requests.get(link, headers=headers).text
  data = BeautifulSoup(html, 'html.parser')
  result = []
  index=1
  while True:
    result += get_list_govuk(df, data)
    # for res in result:
    #     print(res)
    index += 1
    if blogs:
        for i in range(len(blogs)):
            result += read_blog_page(blogs[i][1], blogs[i][0])
            blogs[i][1] = next_blog_page(blogs[i][1], blogs[i][0])
    result = remove_dupes(result)
    if len(result)>=max_results:
      if len(result)>max_results:
        result = result[:(max_results)]
      break
    else:
      link = link[:-1]+str(index)
      html = requests.get(link, headers=headers).text
      data = BeautifulSoup(html, 'html.parser')
  return result

# =============================================================================
# Prints the gathered results and the departments and search terms that the 
# results were based on
# =============================================================================
def print_results(result, refined, search_terms):
    print(f"Here is the {len(result)} result(s) when searching \"{search_terms}\"")
    header = result[0].keys()
    rows = [value.values() for value in result]
    print(tabulate.tabulate(rows, header))

try:
  titles = [elem['Title'] for elem in df]
except:
  titles = list(df['Title'])

# =============================================================================
# Creates the URL for a blog given a unique blog dictionary, search terms,
# and other parameters
# =============================================================================
def create_blog_link(blog, search_terms, sdate=None, edate=None, order_by=None):
  link = blog['Blog Link']+blog['Search']
  if blog['Format']!="+":
    return None
  page = blog['Page']
  if type(page)!=list or len(page)!=2:
    return None
  params = {blog['Keywords']: search_terms}
  if page[0]!='AFTER':
    params[page[0]] = page[1]
  if order_by!=None:
    ordered = order_by.split(" || ")
    params[ordered[0]] = ordered[1] 
  x = requests.get(link, headers=headers, params=params)
  return x.url

# =============================================================================
# Returns list of lists containing a unique blog information dictionary and 
# its corresponding URL
# =============================================================================
def add_blog_links(blogs, search_terms, sdate=None, edate=None, order_by=None):
  links = []
  for blog in blogs:
    URL = create_blog_link(blog, search_terms, sdate, edate, None)#order_by)
    if [blog, URL] not in links:
        links.append([blog, URL])
  return links


# =============================================================================
# Returns number of results from a select URL
# =============================================================================
def get_manual_number(link, reading, search_link=None):
  try:
    x = requests.get(link, headers=headers)
    html = x.text
    soup = BeautifulSoup(html, 'html.parser')
    i = soup.find(reading[0][0], {'class': reading[0][1]})
    links = []
    for j in i.find_all('a'):
      if j['href'] not in links:
        if search_link!=None:
          if j['href'][0]!='#' and not(j['href'].startswith(search_link)):
            links.append(j['href'])
        else:
          links.append(j['href'])
    return len(links)
  except:
    return None

# =============================================================================
# If the page number is inside the URL, this function increments the page number
# by the correct increment number and returns the updated URL
# =============================================================================
def next_page_inner(link, snippet, increment):
  try:
    ind = link.index(snippet)
    first_half = link[:ind]
    sect = link[ind+len(snippet):]
    counter = 1
    number = 0
    while True:
      try:
        number = int(sect[:counter])
        counter += 1
      except:
        counter -= 1
        break
    number = number + increment
    result = first_half + snippet + str(number) + sect[counter:]
    return result
  except:
    return None

# =============================================================================
# Given a link and the corresponding information regarding how to retrieve the information,
#  the tool returns the number of results if the 'get_total_results' calls it
# =============================================================================
def find_blog_number(elem):
    blog = elem[0]
    link = elem[1]
    if blog['Search Link']=='None':
      search_link = None
    else:
      search_link = blog['Search Link']
    number = 0
    reading = []
    find_results = blog['Results'].split("\n")
    for elem in find_results:
      elem = elem.split(" ")
      if len(elem)>2:
        temp = " ".join(elem[0:])
        elem = [elem[0], temp]
      reading.append(elem)
    while True:
      result_num = get_manual_number(link, reading, search_link)
      if result_num==None or result_num==0:
        return number
      number += result_num
      link = next_page_inner(link, 'page/', 1)
    return number

# =============================================================================
# Given a link and the corresponding information regarding how to retrieve the information,
#  the tool updates the link to point to the next page
# =============================================================================
def next_blog_page(link, blog):
    if blog['Page'][0]=='AFTER':
        return next_page_inner(link, blog['Page'][1], blog['Page Increment'])
    l = [-10, -9, -8, -7, -6, -5, -4, -3, -2, -1]
    page = 0
    count = -1
    for i in l:
        try:
            page = int(link[i:])
            count = i
            break
        except:
            pass
    increment = int(blog['Page Increment'])
    return link[:count]+str(page+increment)

# =============================================================================
# Given a link and the corresponding information regarding how to retrieve the information,
#  the tool returns all possible information
# =============================================================================
def read_blog_page(link, blog):
    date="N/A"
    results = []
    conn = requests.get(link, headers=headers)
    
    # If there is no successful connection, return nothing
    if conn.status_code!=200:
        return None
    html = conn.text
    soup = BeautifulSoup(html, 'html.parser')
    find_results = blog['Results'].split("\n")
    first_filter = find_results[0].split(" ")
    if len(first_filter)>2:
        temp = " ".join(first_filter[0:])
        first_filter = [first_filter[0], temp]
    scnd_filter = find_results[1].split(" ")[0]
    # filter html text to find section containing results
    i = soup.find(first_filter[0], {'class': first_filter[1]})
    if i==None:
        return results
    
    # Find any possible dates for the existing features
    dates = []
    if blog['Date']!='None' and len(blog['Date'])!=0:
        if blog['Date']!='None' and len(blog['Date'])!=0:
            d_format = blog['Date'].split(" ")
            print(d_format)
            for d in i.find_all(d_format[0], {'class': d_format[1]}):
                try:
                  found = d.text.strip()
                  d_setup = found.split(" ")
                  if len(d_setup[2])==2:
                    d_setup[2] = "20"+d_setup[2]
                  d_setup[1] = d_setup[1][:3]
                  found = " ".join(d_setup)
                  p = dt.strptime(found, '%d %b %Y')
                  day = ""
                  s_day, s_month, s_year = str(p.day), str(p.month), str(p.year)
                  if len(s_day)==1:
                    day+="0"
                  day+=f"{s_day}/"
                  if len(s_month)==1:
                    day+="0"
                  day+=f"{s_month}/{s_year}"
                except:
                  day = "N/A"
                dates.append(day)
    counter = 0
    page_links = []
    for j in i.find_all(scnd_filter):
        if blog['Search Link']!='None' and len(blog['Search Link'])!=0:
            if j['href'][0]=='#' or j['href'].startswith(blog['Search Link']):
                continue
        if 'http' in j['href'][:6]:
            URL = j['href']
        else:
            if j['href'][0]=="/":
                URL = blog['Blog Link'][:-1] + j['href']
            else:
                URL = blog['Blog Link'] + j['href']
        #if URL in page_links:
        #   continue
        if len(dates)>0:
            try:
                date=dates[counter]
            except:
                date="N/A"
        blog_title = j.text.strip().replace('\n', '')
        blog_title = re.sub(' +', ' ', blog_title)
        entry = {
            'Title': blog_title,
            'URL': URL,
            'Departments, Agencies, and Public bodies': blog['Title'],
            "Abstract": "N/A",
            "Last Updated": "N/A",
            'Date Published': date}
        results.append(entry)
        page_links.append(URL)
        counter += 1
    return results


# =============================================================================
# Adds a hyperlink to a paragraph in a word document
# =============================================================================
def add_hyperlink(paragraph, text, url):
  # This gets access to the document.xml.rels file and gets a new relation id value
  part = paragraph.part
  r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

  # Create the w:hyperlink tag and add needed values
  hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
  hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

  # Create a new run object (a wrapper over a 'w:r' element)
  new_run = docx.text.run.Run(docx.oxml.shared.OxmlElement('w:r'), paragraph)
  new_run.text = text

  # Alternatively, set the run's formatting explicitly
  new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
  new_run.font.underline = True

  # Join all the xml elements together
  hyperlink.append(new_run._element)
  paragraph._p.append(hyperlink)
  return hyperlink

#------------------------ Front-end of Search Tool ----------------------------

# Second page of the tool
def apply_settings():
    global tool_page_num
    tool_page_num = 2
    # ---- Creates list of departments that were selected by the user
    global selected_blogs
    selected_blogs = []
    selected_indices = [index for index, var in enumerate(check_vars) if var.get() == 1]
    chosen_titles = []
    chosen_indices = []
    for index in selected_indices:
      if titles[index]=="All associated agencies and public bodies below":
        #if titles[index-1] not in chosen_titles:
        #  chosen_titles.append(titles[index-1])
        #  chosen_indices.append(index-1)
        counter = 1
        while True:
          if "\t" in titles[index+counter]:
            if titles[index+counter].strip() not in chosen_titles:
              chosen_titles.append(titles[index+counter].strip())
              chosen_indices.append(index+counter)
          else:
            break
          counter += 1
      else:
        if titles[index] not in chosen_titles:
          chosen_titles.append(titles[index])
          chosen_indices.append(index)
    global departments
    departments = chosen_titles
    for elem in departments:
        if elem.strip() in blog_titles:
            print(elem.strip())
            print("HAS REACHABLE BLOG POST")
            blog_index = blog_titles.index(elem.strip())
            selected_blogs.append(all_blogs[blog_index])
            
    selected_data = [full_df[index] for index in chosen_indices]
    
    # ---- Clears the screen
    for widget in root.winfo_children():
        widget.destroy()
    
    # ---- Displays a checklist of existing departments
    # ---- The user selects what departments they want to search
    title_label = tk.Label(root, text="Grey Review\n", bg=main_bg, font=("Arial 42 bold"), fg="#666666") # dark grey (bold as well?????)
    title_label.pack()

    # ---- Prompts user to enter Title/Topic and keywords
    #tk.Label(root, text="Enter the Title/Topic").pack()
    #topic_entry = tk.Entry(root)
    #topic_entry.pack(pady=5)

    keyword_label = tk.Label(root, text="\nEnter keywords", bg=main_bg, font=("Arial 16"))
    keyword_label.pack()
    keyword_entry = tk.Entry(root, highlightbackground=main_bg)
    keyword_entry.pack(pady=5)
    
    value_inside = tk.StringVar()
    value_inside.set("Select an Option")
    options = ["Relevance", "Newest First", "Oldest First"]
    
    sort_label = tk.Label(root, text="\nSelect order of results", bg=main_bg, font=("Arial 16"))
    sort_label.pack()
    sort_menu = tk.OptionMenu(root, value_inside, *options)
    sort_menu.pack(pady=5)

    # ---- The start and end date prompts are only accessible once the 'Enable Advanced Search' is clicked
    sdate_label = tk.Label(root, text="Enter start date (DD/MM/YYY)", bg=main_bg)
    start_date = tk.Entry(root, highlightbackground=main_bg)

    edate_label = tk.Label(root, text="Enter end date (DD/MM/YYY)", bg=main_bg)
    end_date = tk.Entry(root, highlightbackground=main_bg)

    show = False

    # ---- Start and end date prompts appear using this function
    def show_dates():
        global show
        show = True
        sdate_label.pack(before=eas_submit_button)
        start_date.pack(pady=5, before=eas_submit_button)
        edate_label.pack(before=eas_submit_button)
        end_date.pack(pady=5, before=eas_submit_button)
        eas_button.configure(text="Disable Date Range", command=hide_dates)
    
    # ---- Start and end date prompts are hidden using this function
    def hide_dates():
        global show
        show = False
        sdate_label.pack_forget()
        start_date.pack_forget()
        edate_label.pack_forget()
        end_date.pack_forget()
        eas_button.configure(text="Enable Date Range", command=show_dates)

    # Third page of the tool
    def on_submit():
        global tool_page_num
        tool_page_num = 3
        
        # ---- Retrieves information given by user
        global keywords
        global sdate
        global edate
        global sort_by
        
        try:
            keywords = keyword_entry.get()
            sdate = start_date.get()
            edate = end_date.get()
            #topic_value = topic_entry.get()
            sort_val = value_inside.get()
            sort_by = sort_val
        except:
            pass
        
        # Given two dates (format DD/MM/YYYY), it checks if date1 < date2
        def compare_dates(date1, date2):
            date1_l = date1.split("/")
            date2_l = date2.split("/")
            return datetime(int(date1_l[2]), int(date1_l[1]), int(date1_l[0]))<datetime(int(date2_l[2]), int(date2_l[1]), int(date2_l[0]))
        
        # Given a date (string format), checks that the format is correct (DD/MM/YYYY)
        def check_dates(date):
            d_elems = date.split("/")
            try:
                if d_elems[2]=="0000":
                    return None
                dt.date.fromisoformat(f"{d_elems[2]}-{d_elems[1]}-{d_elems[0]}")
                today = dt.date.today().strftime('%d/%m/%Y')
                if not compare_dates(date, today):
                    return None
            except:
                return None
            return date
        
        # ---- Checks if the input is invalid or empty (start and end dates are allowed to be empty)
        
        #if not topic_value:
        #    messagebox.showwarning("Input Error", "Must enter a Title/Topic")
        #    return
        
        if not keywords:
            messagebox.showwarning("Input Error", "Must enter keyword(s)")
            return
        
        if show==True:
            if sdate:
                sdate = check_dates(sdate)
                if sdate==None:
                    print("START DATE PROBLEM")
                    messagebox.showwarning("Input Error", "Please enter the date in the correct format")
                    return
            else:
                sdate = None
            if edate:
                edate = check_dates(edate)
                if edate==None:
                    print("END DATE PROBLEM")
                    messagebox.showwarning("Input Error", "Please enter the date in the correct format")
                    return   
            else:
                edate = None             
            
            if (sdate and edate) and (not compare_dates(sdate, edate)):
                messagebox.showwarning("Input Error", "Please enter sensible start and end dates")
                return
        else:
            sdate, edate = None, None
        # ---- Generates keywords based on the title/topic provided by the user (currently not in use as it does not connect to a generative AI model)
        # keywords = generate_keywords(title, keywords)
        
        # ---- Given the information by the user, the URL is created and used to retrieve the total possible results
        URL = govuk_pubs_link(selected_data, keywords, sdate, edate, sort_by)
        
        global selected_blogs
        global blogs
        if len(selected_blogs)>0:
            print("CREATE BLOG URLS HERE")
            # MAKE BLOG URLS HERE
            blogs = add_blog_links(selected_blogs, keywords, sdate, edate, sort_by)
        else:
            blogs = None
        
        total_results, blogs = get_total_results(URL, blogs)
        
        # ---- If there are no results, user is asked to give new values and try again
        if total_results==None:
            messagebox.showwarning("Value Error", "There were no results given the parameters set, please try different parameters")
            return 
        
        # ---- Cleans the window
        for widget in root.winfo_children():
            widget.destroy()
        
        
        
        # ---- Displays a checklist of existing departments
        # ---- The user selects what departments they want to search
        title_label = tk.Label(root, text="Grey Review\n", bg=main_bg, font=("Arial 42 bold"), fg="#666666") # dark grey (bold as well?????)
        title_label.pack()
        
        # ---- Prompts user to enter the max amount of results they would like given the total number of results
        tk.Label(root, text=f"From the {total_results} result(s) found, how many do you want to be returned?", bg=main_bg, font=("Arial 16")).pack()
        max_results_entry = tk.Entry(root, highlightbackground=main_bg)
        max_results_entry.pack(pady=5)
        
        # Fourth and final page of the tool
        def print_final_result():
            global tool_page_num
            tool_page_num = 4
            # ---- Checks if input given from user for max number of results is invalid
            try:
                max_results = int(max_results_entry.get())
                if max_results<=0 or max_results>total_results:
                    messagebox.showwarning("Input Error", "Please enter a sensible number")
                    return
            except:
                messagebox.showwarning("Input Error", "Please enter a sensible number")
                return
            
            # Exports the selected results to Excel
            def export_results():
                selected_indices = [index for index, var in enumerate(check_vars_in) if var.get() == 1]
                selected_results = [result[index] for index in selected_indices]
                
                # Checks if the select all option was selected
                if len(selected_results)==0:
                    messagebox.showwarning("No results selected", "Please select at least one of the produced results")
                    return
                
                if (len(selected_results)==1 and selected_results[0]=="Select all") or selected_results[0]=="Select all":
                    selected_results = result[1:]
                
                max_results = len(selected_results)
                pd_results = pd.DataFrame(selected_results)
                
                try:
                    pd_results.to_excel('out.xlsx', index=False)
                    messagebox.showinfo("Excel File create successfully", f"New 'out.xlsx' file created at {os.getcwd()} containing {max_results} result(s).")
                except:
                    messagebox.showwarning('Failed to create Excel file', "Could not create a new Excel file")
                return
            
            # Exports the selected results to Word
            def export_word():
                selected_indices = [index for index, var in enumerate(check_vars_in) if var.get() == 1]
                results = [result[index] for index in selected_indices]
                
                if len(results)==0:
                    messagebox.showwarning("No results selected", "Please select at least one of the produced results")
                    return 
                
                if (len(results)==1 and results[0]=="Select all") or results[0]=="Select all":
                    results = result[1:]
                
                max_results = len(results)
                word_result = []
                for i in range(len(results)):
                    word_result.append([results[i]['Departments, Agencies, and Public bodies'], results[i]['Title'], results[i]['URL'], results[i]['Date Published'][-4:], results[i]['Last Updated'][-4:]])
                try:
                    doc = docx.Document()
                    doc.add_heading('Results exported to Word')
                    counter = 1
                    for elem in word_result:
                        p = doc.add_paragraph(f"{counter}. {elem[0]}. ")
                        add_hyperlink(p, elem[1], elem[2])
                        p.add_run(f" ({elem[3][-4:]})")
                        counter += 1
                    doc.save('out.docx')
                    messagebox.showinfo("Word document create successfully", f"New 'out.docx' file created at {os.getcwd()} containing {max_results} result(s).")
                except:
                    messagebox.showwarning('Failed to create Word document', "Could not create a new Word document")
                return
            
            def save_results_file():
                excel_button.pack_forget()
                word_button.pack_forget()
                try:
                    global entered_file
                    entered_file = tk.Entry(root, highlightbackground=main_bg)
                    entered_file.pack(before=back_button)
                    entered_file.delete(0, 'end')
                    entered_file.insert(0, 'Enter filename')
                    # ENTERED FILENAME BECOMES NEW FILENAME
                    # SO IT WOULD BE "YYYY-MM-DD_FILENAME"
                    save_button.configure(text="Save File", command=save_results)
                except:
                    excel_button.pack(before=save_button, pady=5)
                    word_button.pack(before=save_button, pady=5)
                    return
                
            
            # Saves the current results to a new directory to be used later
            def save_results():
                global entered_file
                filename=""
                try:
                    filename = entered_file.get()
                except:
                    messagebox.showwarning('Failed to retrieve filename', "Failed to retrieve filename")
                    return 
                if not os.path.exists(os.getcwd()+"/saved_searches"):
                    os.makedirs(os.getcwd()+"/saved_searches")
                df = pd.DataFrame(result[1:])
                global file_counter
                global keywords
                global sdate
                global edate
                global sort_by
                global departments
                for i in range(len(departments)):
                    departments[i] = departments[i].strip()
                s_departments = " ".join(departments)
                today = dt.date.today().strftime('%Y-%m-%d')
                # all_files = []
                # if len(os.listdir(os.getcwd()+"/saved_searches"))>0:
                #     for name in os.listdir(os.getcwd()+"/saved_searches"):
                #         if name[0]!='.':
                #             if name[:-4] not in all_files:
                #                 all_files.append(name[:-4])
                # if len(all_files)>0:
                #     all_files.reverse()
                #     for name in all_files:
                #         if f"{today}_{file_counter}"==name:
                #             file_counter += 1
                filename = today+" "+filename
                csv_path = os.getcwd()+"/saved_searches/"+filename+".csv"
                txt_path = os.getcwd()+"/saved_searches/"+filename+".txt"
                if os.path.exists(csv_path) or os.path.exists(txt_path):
                    messagebox.showwarning('File already exists', "Filename already exists")
                #filename = f"{today}_{file_counter}"
                try:
                    f = open(os.getcwd()+"/saved_searches/"+filename+".txt", "x")
                    if sort_by=='Select an Option':
                        sort_by = 'Relevance'
                    f.write(f"Sources: {s_departments}\nKeywords: {keywords}\nStart date: {sdate}\nEnd date: {edate}\nSort by: {sort_by}")
                    f.close()
                    df.to_csv(os.getcwd()+"/saved_searches/"+filename+".csv", index=False)
                    file_counter += 1
                    save_button['state'] = 'disabled'
                    messagebox.showinfo("Search saved successfully", f"New '{filename}.txt' and '{filename}.csv' files created at {os.getcwd()}/saved_searches.")
                except:
                    messagebox.showwarning('Failed to save results', "Failed to save results")
                return
            
            # ---- Cleans the window
            for widget in root.winfo_children():
                widget.destroy()
            
            
            
            # ---- Displays a checklist of existing departments
            # ---- The user selects what departments they want to search
            title_label = tk.Label(root, text="Grey Review\n", bg=main_bg, font=("Arial 42 bold"), fg="#666666") # dark grey (bold as well?????)
            title_label.pack()
            
            # ---- Returns list of results matching the information provided by the user
            global blogs
            result = get_pubs(full_df, URL, max_results, blogs)
            top_l = ["Select all"]
            result = top_l + result
            
            tk.Label(root, text="Select the results you would like to be exported", bg=main_bg, font=("Arial 16")).pack()
            tk.Label(root, text="Swipe along to read the full results", bg=main_bg, font=("Arial 13")).pack()
            
            text_inner = ScrolledText(root, width=110, height=25, font=("Arial 13"), cursor="arrow")
            hyperlink = HyperlinkManager(text_inner)
           
            check_vars_in = []
            
            for elem in result:
               var = tk.IntVar()
               check_vars_in.append(var)
               in_text = ""
               if elem=="Select all":
                   in_text = elem
               elif type(elem['Departments, Agencies, and Public bodies'])==list:
                   for dep in elem['Departments, Agencies, and Public bodies']:
                       in_text += dep+", "
                   in_text = in_text.strip()[:-1]
                   elem['Departments, Agencies, and Public bodies'] = in_text
               else:
                   in_text = elem['Departments, Agencies, and Public bodies']
               cb = tk.Checkbutton(text_inner, text=f"{in_text}.", variable=var, anchor='w', bg='white')
               text_inner.window_create('end', window=cb)
               if in_text!="Select all":
                   text_inner.insert(tk.END, f"{elem['Title']}",hyperlink.add(partial(webbrowser.open,elem['URL'])))
                   text_inner.insert(tk.END, f" ({elem['Date Published'][-4:]})")
                   text_inner.insert(tk.END, f" (Last Updated {elem['Last Updated'][-4:]})")
               text_inner.insert('end', '\n')
            text_inner.configure(state='disabled')
            text_inner.pack()
               
            # ---- Exports selected results to Excel
            excel_button = tk.Button(root, text="Export to Excel", command=export_results, highlightbackground=main_bg)
            excel_button.pack(pady=5)
            
            # ---- Exports selected results to Word
            word_button = tk.Button(root, text="Export to Word", command=export_word, highlightbackground=main_bg)
            word_button.pack(pady=5)
            
            save_button = tk.Button(root, text="Save Current Results", command=save_results_file, highlightbackground=main_bg)
            save_button.pack(pady=5)
            
            # ---- Sends user back to third page of tool
            back_button = tk.Button(root, text="Back", command=on_submit, highlightbackground=main_bg)
            back_button.pack(pady=5)
            
            # ---- Sends user to beginning
            back_start_button = tk.Button(root, text="Back to Start", command=front_page, highlightbackground=main_bg)
            back_start_button.pack(pady=5)
            
            help_button = tk.Button(root, text="Help", command=help_page, highlightbackground=main_bg)
            help_button.place(rely=1.0, relx=1.0, x=0, y=0, anchor=tk.SE)
        
        # ---- Sends new information to the show results page
        submit_button = tk.Button(root, text="Submit", command=print_final_result, highlightbackground=main_bg)
        submit_button.pack(pady=10)
        
        # ---- Sends user back to second page of tool
        back_button = tk.Button(root, text="Back", command=apply_settings, highlightbackground=main_bg)
        back_button.pack(pady=10)
        
        help_button = tk.Button(root, text="Help", command=help_page, highlightbackground=main_bg)
        help_button.place(rely=1.0, relx=1.0, x=0, y=0, anchor=tk.SE)
    
    # ---- Click to show start and end date prompts
    eas_button = tk.Button(root, text="Enable Date Range", command=show_dates, highlightbackground=main_bg)
    eas_button.pack(pady=10)

    # ---- Sends new information to retrieve information from gov.uk
    eas_submit_button = tk.Button(root, text="Submit", command=on_submit, highlightbackground=main_bg)
    eas_submit_button.pack(pady=10)
    
    # ---- Sends user back to first page of tool
    back_button = tk.Button(root, text="Back", command=front_page, highlightbackground=main_bg)
    back_button.pack(pady=10)
    
    help_button = tk.Button(root, text="Help", command=help_page, highlightbackground=main_bg)
    help_button.place(rely=1.0, relx=1.0, x=0, y=0, anchor=tk.SE)

# Shows all saved searches
def use_saved():
    global tool_page_num
    tool_page_num = 5
    if os.path.exists(os.getcwd()+"/saved_searches"):
        
        # Retrieves all saved files
        all_files = os.listdir(os.getcwd()+"/saved_searches")
        result = []
        files = []
        for name in all_files:
            if name[0]!='.':
                filename = name.split(".")[0]
                if filename+".txt" in all_files and filename+".csv" in all_files and filename not in files:
                    files.append(filename)
                    saved_file = {}
                    saved_file['filename'] = filename
                    # all_files.remove(filename+".txt")
                    # all_files.remove(filename+".csv")
                    """f = open(os.getcwd()+"/saved_searches/"+filename+".txt", "r")
                    content = f.read().split("\n")
                    f.close()
                    for elem in content:
                        pair = elem.split(":")
                        saved_file[pair[0]] = pair[1]"""
                    with open(os.getcwd()+"/saved_searches/"+filename+".csv") as f:
                        dataset = [{k: v for k, v in row.items()} for row in csv.DictReader(f, skipinitialspace=True)]
                    saved_file['Data'] = dataset
                    result.append(saved_file)
        global saved_results
        saved_results = result
        for widget in root.winfo_children():
            widget.destroy()
        
        
        
        # ---- Displays a checklist of existing departments
        # ---- The user selects what departments they want to search
        title_label = tk.Label(root, text="Grey Review\n", bg=main_bg, font=("Arial 42 bold"), fg="#666666") # dark grey (bold as well?????)
        title_label.pack()
        
        saved_check_vars = []
        
        # ---- Displays a checklist of existing saved results
        # ---- The user selects what departments they want to search
        tk.Label(root, text="Select only one save file to load\n", bg=main_bg, font=("Arial 16")).pack()
        
        text = ScrolledText(root, width=90, height=32, cursor="arrow")
        text.pack(fill=tk.Y)
        
        for data in result:
            # saved_keys = data['Keywords']
            # s_date = data['Start date']
            # e_date = data['End date']
            # sort = data['Sort by']
            # if sort=="Select an Option":
            #     sort = "Relevance"
            # s_departments = data['Sources']
            # s_deps = ""
            # for dep in s_departments.split("||"):
            #     s_deps += dep+", "
            # s_deps = s_deps.strip()[:-1]
            # if len(saved_keys)>10:
            #     saved_keys = saved_keys[:10]+"..."
            # if s_date=="None":
            #     s_date = "N/A"
            # if e_date=="None":
            #     e_date ="N/A"
            background = "white"
            var = tk.IntVar()
            saved_check_vars.append(var)
            cb = tk.Checkbutton(text, text=f"{data['filename']}", variable=var, bg=background, anchor='w')
            text.window_create('end', window=cb)
            text.insert('end', '\n')
        
        text.configure(state=tk.DISABLED)
    else:
        messagebox.showwarning('No saved searches directory', f"No saved searches directory exists at '{os.getcwd()}'")
        return

    # Deletes a selected saved file
    def delete_save():
        selected_indices = [index for index, var in enumerate(saved_check_vars) if var.get() == 1]
        if len(selected_indices)!=1:
            messagebox.showwarning('Error', "Please only select one option")
            return
        global saved_results
        selected_results = [saved_results[index] for index in selected_indices][0]
        try:
            os.remove(os.getcwd()+"/saved_searches/"+selected_results['filename']+".txt")
            os.remove(os.getcwd()+"/saved_searches/"+selected_results['filename']+".csv")
            use_saved()
        except:
            messagebox.showwarning("Could not delete file", "Could not delete saved data")
            return
    
    def present_saved():
        global tool_page_num
        tool_page_num = 6
        
        selected_indices = [index for index, var in enumerate(saved_check_vars) if var.get() == 1]
        if len(selected_indices)!=1:
            messagebox.showwarning('Error', "Please only select one option")
            return
        global saved_results
        selected_results = [saved_results[index] for index in selected_indices][0]
        
        for widget in root.winfo_children():
            widget.destroy()
        
        
        
        
        # ---- Displays a checklist of existing departments
        # ---- The user selects what departments they want to search
        title_label = tk.Label(root, text="Grey Review", bg=main_bg, font=("Arial 42 bold"), fg="#666666") # dark grey (bold as well?????)
        title_label.pack()
        
        top_l = ["Select all"]
        result = top_l + selected_results['Data']
        
        tk.Label(root, text="Select the results you would like to be exported\n", bg=main_bg, font=("Arial 16")).pack()
        tk.Label(root, text="Swipe along to read the full results\n", bg=main_bg, font=("Arial 16")).pack()
        
        text_inner = ScrolledText(root, width=110, height=25, font=("Arial 13"), cursor="arrow")
        hyperlink = HyperlinkManager(text_inner)
       
        check_vars_in = []
        
        # Presnets all results from the saved file as a checkbutton list
        for elem in result:
           var = tk.IntVar()
           check_vars_in.append(var)
           in_text = ""
           if elem=="Select all":
               in_text = elem
           elif type(elem['Departments, Agencies, and Public bodies'])==list:
               for dep in elem['Departments, Agencies, and Public bodies']:
                   in_text += dep+", "
               in_text = in_text.strip()[:-1]
               elem['Departments, Agencies, and Public bodies'] = in_text
           else:
               in_text = elem['Departments, Agencies, and Public bodies']
           cb = tk.Checkbutton(text_inner, text=f"{in_text}.", variable=var, anchor='w', bg='white')
           text_inner.window_create('end', window=cb)
           if in_text!="Select all":
               text_inner.insert(tk.END, f"{elem['Title']}",hyperlink.add(partial(webbrowser.open,elem['URL'])))
               text_inner.insert(tk.END, f" ({elem['Date Published'][-4:]})")
               text_inner.insert(tk.END, f" (Last Updated {elem['Last Updated'][-4:]})")
           text_inner.insert('end', '\n')
        text_inner.configure(state='disabled')
        text_inner.pack()
        
        # Exports selected results to Excel
        def export_results():
            selected_indices = [index for index, var in enumerate(check_vars_in) if var.get() == 1]
            selected_results = [result[index] for index in selected_indices]
            
            # Checks if the select all option was selected
            if len(selected_results)==0:
                messagebox.showwarning("No results selected", "Please select at least one of the produced results")
                return
            
            if (len(selected_results)==1 and selected_results[0]=="Select all") or selected_results[0]=="Select all":
                selected_results = result[1:]
            
            max_results = len(selected_results)
            pd_results = pd.DataFrame(selected_results)
            
            try:
                pd_results.to_excel('out.xlsx', index=False)
                messagebox.showinfo("Excel File create successfully", f"New 'out.xlsx' file created at {os.getcwd()} containing {max_results} result(s).")
            except:
                messagebox.showwarning('Failed to create Excel file', "Could not create a new Excel file")
            return
        
        # Exports the selected results to Word
        def export_word():
            selected_indices = [index for index, var in enumerate(check_vars_in) if var.get() == 1]
            results = [result[index] for index in selected_indices]
            
            if len(results)==0:
                messagebox.showwarning("No results selected", "Please select at least one of the produced results")
            
            if (len(results)==1 and results[0]=="Select all") or results[0]=="Select all":
                results = result[1:]
                return
            
            max_results = len(results)
            word_result = []
            for i in range(len(results)):
                word_result.append([results[i]['Departments, Agencies, and Public bodies'], results[i]['Title'], results[i]['URL'], results[i]['Date Published'][-4:], results[i]['Last Updated'][-4:]])
            try:
                doc = docx.Document()
                doc.add_heading('Results exported to Word')
                counter = 1
                for elem in word_result:
                    p = doc.add_paragraph(f"{counter}. {elem[0]}. ")
                    add_hyperlink(p, elem[1], elem[2])
                    p.add_run(f" ({elem[3][-4:]})")
                    p.add_run(f" (Last Updated {elem[4][-4:]})")
                    counter += 1
                doc.save('out.docx')
                messagebox.showinfo("Word document create successfully", f"New 'out.docx' file created at {os.getcwd()} containing {max_results} result(s).")
            except:
                messagebox.showwarning('Failed to create Word document', "Could not create a new Word document")
            return
           
        # ---- Exports selected results to Excel
        submit_button = tk.Button(root, text="Export to Excel", command=export_results, highlightbackground=main_bg)
        submit_button.pack(pady=10)
        
        # ---- Exports selected results to Word
        submit_button = tk.Button(root, text="Export to Word", command=export_word, highlightbackground=main_bg)
        submit_button.pack(pady=10)
        
        back_button = tk.Button(root, text="Back", command=use_saved, highlightbackground=main_bg)
        back_button.pack(pady=10)
        
        help_button = tk.Button(root, text="Help", command=help_page, highlightbackground=main_bg)
        help_button.place(rely=1.0, relx=1.0, x=0, y=0, anchor=tk.SE)
        
        return
    
    submit_button = tk.Button(root, text="Submit", command=present_saved, highlightbackground=main_bg)
    submit_button.pack(pady=10)
    
    del_button = tk.Button(root, text="Delete", command=delete_save, highlightbackground=main_bg)
    del_button.pack(pady=10)
    
    # ---- Sends user back to first page of tool
    back_button = tk.Button(root, text="Back", command=front_page, highlightbackground=main_bg)
    back_button.pack(pady=10)
    
    help_button = tk.Button(root, text="Help", command=help_page, highlightbackground=main_bg)
    help_button.place(rely=1.0, relx=1.0, x=0, y=0, anchor=tk.SE)
    
    return


data = df

entered_file = None
departments = None
keywords = None
sdate= None
edate = None
sort_by = None
file_counter = 1
saved_results = None
selected_blogs = None
blogs = None
tool_page_num = 0



titles = []
full_df = []
for elem in df:
    titles.append(elem['Title'])
    full_df.append(elem)
    if elem['Works with']!='None':
        titles.append("All associated agencies and public bodies below")
        for worker in elem['Works with']:
            titles.append(f"\t{worker['Title']}")
            full_df.append(worker)

root = tk.Tk()
root.geometry("1000x850")
root.title("Grey Literature Search Tool")
main_bg = "#d9e2f3"
root.configure(bg=main_bg)

frame = tk.Frame(root)
frame.pack(padx=10, pady=10)

check_vars = []

# If the help button is clicked, a corresponding help message pops up on screen
def help_page():
    text = "Help: "
    global tool_page_num
    if tool_page_num == 1:
        text += "This page is for you to select which agencies and public bodies you would like to retrieve information from"
    elif tool_page_num == 2:
        text += "This page is for you to:\n- Enter keywords\n- Select how you want the results ordered (default set to relevance)\n- Enter a date range if applicable"
    elif tool_page_num == 3:
        text += "This page is for you to enter the max number of results that you would like to see on the next page"
    elif tool_page_num == 4:
        text += "This page is for you to view the results and to either:\n- Export a selected number of results to Excel or Word\n- Save the current results"
    elif tool_page_num == 5:
        text += "This page is for you to view existing save files to view. Select one file to view the results"
    elif tool_page_num == 6:
        text += "This page is for you to view the saved results and to export a selected number of results to either Excel or Word"
    messagebox.showinfo("Help Message", text)
    return 


def front_page():
    global tool_page_num
    tool_page_num = 1
    
    global check_vars
    check_vars = []
    
    for widget in root.winfo_children():
        widget.destroy()
    
    
    
    # ---- Displays a checklist of existing departments
    # ---- The user selects what departments they want to search
    title_label = tk.Label(root, text="Grey Review\n", bg=main_bg, font=("Arial 42 bold"), fg="#666666") # dark grey (bold as well?????)
    title_label.pack()
    
    tk.Label(root, text="Select the departments, agencies, and public bodies you want to search from\n", bg=main_bg, font=("Arial 16")).pack()
    
    text = ScrolledText(root, width=90, height=32, cursor="arrow")
    text.pack(fill=tk.Y)
    
    for title in titles:
        background = "white"
        if title=="All associated agencies and public bodies below":
            background = "light grey"
        var = tk.IntVar()
        check_vars.append(var)
        if "\t" not in title and title!="All associated agencies and public bodies below":
            cb = tk.Checkbutton(text, text=title, variable=var, bg=background, anchor='w', font=("Arial", "14", "bold"))
        else:
            cb = tk.Checkbutton(text, text=title, variable=var, bg=background, anchor='w')
        text.window_create('end', window=cb)
        text.insert('end', '\n')
    
    text.configure(state=tk.DISABLED)
    
    # ---- Once they have selected the departments they want, they click the button
    # ----  and go to the second page of the tool
    print_button = tk.Button(root, text="Proceed to Searching", command=apply_settings, highlightbackground=main_bg)
    print_button.pack(pady=10)
    
    print_button = tk.Button(root, text="Use Saved Search", command=use_saved, highlightbackground=main_bg)
    print_button.pack(pady=10)
    
    help_button = tk.Button(root, text="Help", command=help_page, highlightbackground=main_bg)
    help_button.place(rely=1.0, relx=1.0, x=0, y=0, anchor=tk.SE)

front_page()

# Run the application
root.mainloop()